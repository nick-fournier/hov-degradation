import pandas as pd
import os

from datetime import date
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.shared import Inches


class PlotsToDocx:

    def __init__(self, path, start_date, end_date):
        self.path = path
        self.dates = start_date + '_to_' + end_date
        self.agg_results = self.aggregate_results()
        self.docx = self.img_to_doc()

        self.docx.save(self.path + '/results/HOV Deg Results Draft_' + str(date.today()) + '.docx')



    def aggregate_results(self):
        df_meta = pd.read_csv(self.path + "data/meta_2020-11-16.csv")
        df_pred = pd.read_csv(self.path + 'results/predictions_D7_' + self.dates + '.csv')

        total = df_meta.loc[df_meta.Type == 'HV'].ID.count()
        analyzed = df_pred.iloc[:, 0].count()
        pred_sup = df_pred.loc[df_pred.preds_classification > 0].iloc[:, 0].count()
        pred_unsup = df_pred.loc[df_pred.preds_unsupervised > 0].iloc[:, 0].count()

        res = {'Total HOVs': total,
               'Analyzed HOVs': analyzed,
               'Identified Misconfigurations (unsupervised)': pred_unsup,
               'Identified Misconfigurations (supervised)': pred_sup,
               'Analysis date': start_date + " to " + end_date}

        return res

    def img_to_doc(self):
        # Set up the empty doc
        doc = Document()

        plot_path = 'results/misconfig_plots_'
        strip_path = 'results/strip_maps/'

        para = doc.add_paragraph()
        run = para.add_run()
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(18)
        font.bold = True
        for key in self.agg_results.keys():
            run.add_text( key + ": " + str(self.agg_results[key]))
            run.add_break()
        run.add_break(WD_BREAK.PAGE)

        for id_dir in os.listdir(doc_path + plot_path + self.dates):
            doc.add_heading('Sensor: ' + id_dir, level=2)
            para = doc.add_paragraph()
            run = para.add_run()
            font.name = 'Calibri'
            font = run.font
            # run.add_break()
            for img in os.listdir(doc_path + plot_path + self.dates + '/' + id_dir):
                run.add_picture(doc_path + plot_path + self.dates + '/' + id_dir + '/' + img, height=Inches(3))
                # run.add_break()
            run.add_break()
            run.add_picture(doc_path + strip_path + id_dir + "_strip.png", width=Inches(6))
            run.add_text('Comments:')
            run.add_break()
            run.add_break(WD_BREAK.PAGE)

        return doc

if __name__ == '__main__':
    doc_path = '../../experiments/district_7/'
    # doc_path = 'experiments/district_7/'
    start_date = '2020-12-06'
    end_date = '2020-12-12'

    document = PlotsToDocx(doc_path, start_date, end_date)


