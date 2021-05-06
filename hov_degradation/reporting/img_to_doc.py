"""
Copyright ©2021. The Regents of the University of California (Regents). All Rights Reserved. Permission to use, copy, modify, and distribute this software and its documentation for educational, research, and not-for-profit purposes, without fee and without a signed licensing agreement, is hereby granted, provided that the above copyright notice, this paragraph and the following two paragraphs appear in all copies, modifications, and distributions. Contact The Office of Technology Licensing, UC Berkeley, 2150 Shattuck Avenue, Suite 510, Berkeley, CA 94720-1620, (510) 643-7201, otl@berkeley.edu, http://ipira.berkeley.edu/industry-info for commercial licensing opportunities.

IN NO EVENT SHALL REGENTS BE LIABLE TO ANY PARTY FOR DIRECT, INDIRECT, SPECIAL, INCIDENTAL, OR CONSEQUENTIAL DAMAGES, INCLUDING LOST PROFITS, ARISING OUT OF THE USE OF THIS SOFTWARE AND ITS DOCUMENTATION, EVEN IF REGENTS HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.

REGENTS SPECIFICALLY DISCLAIMS ANY WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE. THE SOFTWARE AND ACCOMPANYING DOCUMENTATION, IF ANY, PROVIDED HEREUNDER IS PROVIDED "AS IS". REGENTS HAS NO OBLIGATION TO PROVIDE MAINTENANCE, SUPPORT, UPDATES, ENHANCEMENTS, OR MODIFICATIONS.
"""

import pandas as pd
import os
import json

from datetime import date
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.shared import Inches


def check_before_reading(path):
    with open(path) as f:
        first_line = f.readline()

    if "," in first_line:
        return pd.read_csv(path)

    if "\t" in first_line:
        return pd.read_csv(path, sep="\t")

class PlotsToDocx:
    def __init__(self, inpath, outpath, plot_date, date_range_string):
        # Amend string
        if inpath[-1] is not "/":
            self.inpath = inpath + "/"
        else:
            self.inpath = inpath

        # Amend string
        if outpath[-1] is not "/":
            self.outpath = outpath + "/"
        else:
            self.outpath = outpath

        # Read meta data
        self.flist = pd.Series(os.listdir(self.inpath))
        file = self.flist[self.flist.str.contains("meta")][0]
        self.df_meta = check_before_reading(self.inpath + file)
        self.district = str(self.df_meta.loc[0, 'District'])

        # Run
        self.dates = date_range_string
        self.plot_date = plot_date
        self.agg_results = self.aggregate_results()
        self.mis_id_text, self.neighbors, self.reconfig_ids = self.get_labels()



        self.doc = self.img_to_doc()

    def get_labels(self):
       # Load JSON files
        with open(self.outpath + "processed data/D" + self.district + "_neighbors_" + self.dates + ".json") as f:
            neighbors = json.load(f)

        with open(self.outpath + "misconfigs_ids_D" + self.district + "_" + self.dates + ".json") as f:
            mis_ids = json.load(f)

        if os.path.isfile(self.outpath + "fixed_sensors.json"):
            with open(self.outpath + "fixed_sensors.json") as f:
                reconfig_ids = json.load(f)
        else:
            reconfig_ids = {}

        # Get unique predictions
        mis_ids_unique = mis_ids['classification'] + mis_ids['unsupervised']
        mis_ids_unique = pd.Series(mis_ids_unique).sort_values().unique().tolist()

        # Create text dict
        mis_id_text = {}
        for id in mis_ids_unique:
            method = [m for m in ['classification', 'unsupervised'] if id in mis_ids[m]]
            if len(method) > 1:
                method = 'both ' + ' and '.join(method) + ' methods'
            else:
                method += ' method only'
                method = ''.join(method)
            mis_id_text[id] = 'Potential misconfiguration detected by ' + method + '. '

        return mis_id_text, neighbors, reconfig_ids

    def aggregate_results(self):
        # Meta data
        flist = pd.Series(os.listdir(self.outpath))
        f = list(flist[flist.str.contains("meta")])[0]
        df_meta = pd.read_csv(self.outpath + f)
        df_pred = pd.read_csv(self.outpath + 'predictions_D' + self.district + '_' + self.dates + '.csv')

        total = df_meta.loc[df_meta.Type == 'HV'].ID.count()
        analyzed = df_pred.iloc[:, 0].count()
        pred_sup = df_pred.loc[df_pred.preds_classification > 0].iloc[:, 0].count()
        pred_unsup = df_pred.loc[df_pred.preds_unsupervised > 0].iloc[:, 0].count()

        res = {'Total HOVs': total,
               'Analyzed HOVs': analyzed,
               'Identified Misconfigurations (unsupervised)': pred_unsup,
               'Identified Misconfigurations (supervised)': pred_sup,
               'Analysis date': self.dates.replace("_", " ")}

        return res

    def img_to_doc(self):
        # Set up the empty doc
        doc = Document()

        plot_path = 'plots_misconfigs_'
        strip_path = ''
        i = 0
        dir = self.outpath
        delim = "\\" if "\\" in self.outpath else "/"
        while dir:
            if os.path.isdir(dir) and "strip_maps" in os.listdir(dir):
                strip_path = dir + "/strip_maps/"
                break
            i += 1
            dir = "/".join(self.outpath.split(delim)[:-i])

        date_string = pd.to_datetime(self.plot_date).day_name() + ' ' + self.plot_date

        para = doc.add_paragraph()
        run = para.add_run()
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(18)
        font.bold = True

        for key in self.agg_results.keys():
            run.add_text( key + ": " + str(self.agg_results[key]))
            run.add_break()
        run.add_break(WD_BREAK.PAGE)

        for id_dir in os.listdir(self.outpath + plot_path + self.dates):
            figpath = self.outpath + plot_path + self.dates + '/' + id_dir + '/'

            doc.add_heading('Sensor: ' + id_dir, level=2)
            para = doc.add_paragraph()
            run = para.add_run()
            font.name = 'Calibri'
            font = run.font
            run.add_text(self.mis_id_text[int(id_dir)] + 'Plotted using data for {}.'.format(date_string))
            run.add_break()

            # for img in os.listdir(doc_path + plot_path + self.dates + '/' + id_dir):
            #     run.add_picture(doc_path + plot_path + self.dates + '/' + id_dir + '/' + img, height=Inches(3))
            #     # run.add_break()

            run.add_picture(figpath + id_dir + '_lat.png', height=Inches(2.35))
            run.add_picture(figpath + id_dir + '_long.png', height=Inches(2.35))

            if id_dir in list(self.reconfig_ids.keys()):
                run.add_picture(figpath + id_dir + '_fix.png', width=Inches(6))

            run.add_break()
            if os.path.isfile(strip_path + id_dir + "_strip.png"):
                run.add_picture(strip_path + id_dir + "_strip.png", width=Inches(6))
            run.add_text('Comments:')
            run.add_break()
            run.add_break(WD_BREAK.PAGE)

        return doc

    def save(self):
        self.doc.save(self.outpath + '/HOV plots_' + str(date.today()) + '.docx')



